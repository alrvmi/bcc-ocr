# prod.py — универсальный финальный скрипт
import os
from pdf2image import convert_from_path
from paddleocr import PaddleOCR
from docx import Document

# ---------- Настройки ----------
SCANS_DIR = os.path.join("data", "scans")   # входные PDF
BASE_OUTPUT_DIR = os.path.join("data", "output")
POPPLER_PATH = r"C:\poppler\Library\bin"    # если poppler установлен, иначе None
DPI = 300
CONF_THRESHOLD = None  # None = не фильтровать по confidence, или float e.g. 0.5

# ---------- Утилиты ----------
def extract_lines(result):
    """
    Универсально извлекает (text, score) из результата PaddleOCR,
    поддерживает разные форматы, которые мы наблюдали.
    Возвращает список кортежей (text:str, score:float|None)
    """
    lines = []
    if not result:
        return lines

    # 1) Если словарь (редкий случай)
    if isinstance(result, dict):
        rec_texts = result.get("rec_texts", [])
        rec_scores = result.get("rec_scores", [])
        for i, t in enumerate(rec_texts):
            s = rec_scores[i] if i < len(rec_scores) else None
            lines.append((str(t), s))
        return lines

    # 2) Если список
    if isinstance(result, (list, tuple)):
        for elem in result:
            # elem может быть dict (внутри списка) с rec_texts
            if isinstance(elem, dict):
                rec_texts = elem.get("rec_texts", [])
                rec_scores = elem.get("rec_scores", [])
                for i, t in enumerate(rec_texts):
                    s = rec_scores[i] if i < len(rec_scores) else None
                    lines.append((str(t), s))
                continue

            # elem может быть "page" — список строк
            if isinstance(elem, (list, tuple)):
                for line in elem:
                    # line обычно выглядит как [box, (text, score)]
                    try:
                        if isinstance(line, (list, tuple)) and len(line) == 2:
                            second = line[1]
                            # second может быть tuple (text, score) или прямо строка
                            if isinstance(second, (list, tuple)):
                                text = second[0]
                                score = second[1] if len(second) > 1 else None
                                lines.append((str(text), score))
                            elif isinstance(second, dict):
                                # fallback, иногда могут быть dict поля
                                text = second.get("rec_text") or second.get("text") or str(second)
                                score = second.get("rec_score") or second.get("score")
                                lines.append((str(text), score))
                            else:
                                # простая строка
                                lines.append((str(second), None))
                        elif isinstance(line, str):
                            lines.append((line, None))
                        else:
                            # придуманный формат — stringfy
                            lines.append((str(line), None))
                    except Exception:
                        # безопасный fallback: строковое представление
                        try:
                            lines.append((str(line), None))
                        except:
                            continue
                continue

            # если элемент простой текст
            if isinstance(elem, str):
                lines.append((elem, None))
            else:
                # fallback
                lines.append((str(elem), None))

    return lines

# ---------- Основной процесс ----------
def process_pdf(pdf_path, poppler_path=None):
    basename = os.path.splitext(os.path.basename(pdf_path))[0]
    out_dir = os.path.join(BASE_OUTPUT_DIR, basename)
    os.makedirs(out_dir, exist_ok=True)

    print(f"\n[PROCESS] {pdf_path} → {out_dir}")

    # конвертация pdf -> страницы
    try:
        if poppler_path and os.path.exists(poppler_path):
            pages = convert_from_path(pdf_path, dpi=DPI, poppler_path=poppler_path)
        else:
            pages = convert_from_path(pdf_path, dpi=DPI)
    except Exception as e:
        print(f"[ERROR] Не удалось конвертировать {pdf_path}: {e}")
        return

    print(f"[OK] Конвертировано {len(pages)} страниц")

    image_paths = []
    for i, page in enumerate(pages, start=1):
        img_path = os.path.join(out_dir, f"page_{i}.png")
        page.save(img_path, "PNG")
        image_paths.append(img_path)
        print(f"  [SAVED] {img_path}")

    # Инициализируем OCR один раз
    # используем современный параметр, если доступен
    try:
        ocr = PaddleOCR(lang='ru', use_textline_orientation=True)
    except TypeError:
        # fallback на старый параметр
        ocr = PaddleOCR(lang='ru', use_angle_cls=False)

    # создаём doc и txt
    doc = Document()
    txt_path = os.path.join(out_dir, "result.txt")

    with open(txt_path, "w", encoding="utf-8") as txt_file:
        for page_idx, img_path in enumerate(image_paths, start=1):
            print(f"[OCR] Обрабатываю {img_path} (страница {page_idx})")
            # вызываем predict если доступен
            try:
                if hasattr(ocr, "predict"):
                    raw = ocr.predict(img_path)
                else:
                    raw = ocr.ocr(img_path)
            except Exception as e:
                print(f"[ERROR] OCR упал для {img_path}: {e}")
                continue

            # извлекаем пары (text, score)
            pairs = extract_lines(raw)
            # фильтруем пустые строки
            pairs = [(t.strip(), s) for (t, s) in pairs if isinstance(t, str) and t.strip()]

            # запись заголовка страницы
            doc.add_heading(f"{basename} — Страница {page_idx}", level=2)
            txt_file.write(f"--- Страница {page_idx} ---\n")

            if not pairs:
                doc.add_paragraph("[Пусто или нераспознано]")
                txt_file.write("[Пусто или нераспознано]\n\n")
                print(f"  [WARN] Нет строк для {img_path}")
                continue

            for text, score in pairs:
                # можно включить порог уверенности: if CONF_THRESHOLD and score and score < CONF_THRESHOLD: skip
                if CONF_THRESHOLD is not None and score is not None:
                    try:
                        if float(score) < float(CONF_THRESHOLD):
                            # помечаем, но не добавляем — тут можно изменить логику
                            doc.add_paragraph(f"{text}  (low_conf={score:.2f})")
                            txt_file.write(f"{text}  (low_conf={score:.2f})\n")
                            continue
                    except Exception:
                        pass

                # запись в doc & txt
                if score is not None:
                    doc.add_paragraph(f"{text}  (conf={score:.2f})")
                    txt_file.write(f"{text}  (conf={score:.2f})\n")
                else:
                    doc.add_paragraph(text)
                    txt_file.write(f"{text}\n")

            txt_file.write("\n")

    # сохраняем docx
    docx_path = os.path.join(out_dir, "result.docx")
    doc.save(docx_path)
    print(f"[DONE] Сохранены: {txt_path} и {docx_path}")

# ---------- Запуск для всех PDF в папке scans ----------
if __name__ == "__main__":
    if not os.path.isdir(SCANS_DIR):
        print(f"[FATAL] Папка со сканами не найдена: {SCANS_DIR}")
        raise SystemExit(1)

    pdf_files = [f for f in os.listdir(SCANS_DIR) if f.lower().endswith(".pdf")]
    if not pdf_files:
        print(f"[FATAL] PDF-файлы не найдены в {SCANS_DIR}")
        raise SystemExit(1)

    # проверяем poppler путь (если он указан)
    poppler_path = POPPLER_PATH if POPPLER_PATH and os.path.exists(POPPLER_PATH) else None
    if POPPLER_PATH and not poppler_path:
        print(f"[WARN] Указанный POPPLER_PATH '{POPPLER_PATH}' не найден — будет использован PATH или системный poppler (если есть).")

    for pdf in pdf_files:
        pdf_path = os.path.join(SCANS_DIR, pdf)
        process_pdf(pdf_path, poppler_path=poppler_path)
